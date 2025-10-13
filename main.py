import os
from dotenv import load_dotenv
import io
import json
import re
import tempfile
from pathlib import Path
from typing import List, Dict, Optional
from datetime import datetime
from mermaid import Mermaid

import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from sentence_transformers import SentenceTransformer
import google.generativeai as genai
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
    
from graphviz import Digraph
import matplotlib.pyplot as plt
import networkx as nx


# 👇 Add Graphviz bin path manually
os.environ["PATH"] += os.pathsep + r"C:\Program Files\Graphviz\bin"

# Load environment variables
load_dotenv()

# ---------------- CONFIG ----------------
EMBED_MODEL_NAME = os.getenv("EMBED_MODEL_NAME", "all-MiniLM-L6-v2")
GEMINI_MODEL = os.getenv("GEMINI_MODEL", "gemini-2.0-flash")
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")

# ---------------- INIT ----------------
st.set_page_config(page_title="BRD Automation Studio", layout="wide", page_icon="🤖")

st.markdown("""
<style>
    .main-header {
        font-size: 2.5rem;
        font-weight: bold;
        color: #1f77b4;
        text-align: center;
        margin-bottom: 1rem;
    }
    .stButton>button {
        width: 100%;
    }
</style>
""", unsafe_allow_html=True)

st.markdown('<h1 class="main-header">🤖 BRD Automation Studio v3.0</h1>', unsafe_allow_html=True)

if not GOOGLE_API_KEY:
    st.error("⚠️ GOOGLE_API_KEY not set. Please set it in your .env file.")
    st.stop()

# Configure Gemini
try:
    genai.configure(api_key=GOOGLE_API_KEY)
    st.success(f"✅ Connected to Gemini API - Model: {GEMINI_MODEL}")
except Exception as e:
    st.error(f"❌ Failed to configure Gemini: {e}")
    st.stop()

# Initialize session state
if 'vector_store' not in st.session_state:
    st.session_state.vector_store = {'vectors': []}
if 'sample_content' not in st.session_state:
    st.session_state.sample_content = None
if 'video_analysis' not in st.session_state:
    st.session_state.video_analysis = None
if 'process_file_content' not in st.session_state:
    st.session_state.process_file_content = None

# ---------------- VECTOR STORE ----------------
class SimpleVectorStore:
    def __init__(self):
        self.vectors = []
        
    def upsert(self, vectors):
        self.vectors.extend(vectors)
        
    def query(self, vector, top_k=6):
        if not self.vectors:
            return {"matches": []}
        
        scores = []
        for vid, vec, metadata in self.vectors:
            try:
                score = sum(a * b for a, b in zip(vector, vec)) / (
                    (sum(a * a for a in vector) ** 0.5) * 
                    (sum(b * b for b in vec) ** 0.5)
                )
                scores.append((score, vid, metadata))
            except:
                continue
        
        scores.sort(reverse=True)
        matches = []
        for score, vid, metadata in scores[:top_k]:
            matches.append({
                "id": vid,
                "score": score,
                "metadata": metadata
            })
        return {"matches": matches}
    
    def clear(self):
        self.vectors = []

vector_store = SimpleVectorStore()
if 'vector_store' in st.session_state and 'vectors' in st.session_state.vector_store:
    vector_store.vectors = st.session_state.vector_store['vectors']

@st.cache_resource
def load_embedding_model(name: str = EMBED_MODEL_NAME):
    try:
        return SentenceTransformer(name)
    except Exception as e:
        st.error(f"Failed to load embedding model: {e}")
        return None

embed_model = load_embedding_model()

# ---------------- VIDEO PROCESSING ----------------
def extract_video_frames(video_bytes, num_frames=8):
    try:
        import cv2
        from PIL import Image
        import numpy as np
        
        tfile = tempfile.NamedTemporaryFile(delete=False, suffix='.mp4')
        tfile.write(video_bytes)
        tfile.close()
        
        cap = cv2.VideoCapture(tfile.name)
        total_frames = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
        
        if total_frames == 0:
            os.unlink(tfile.name)
            return []
        
        frames = []
        frame_indices = [int(total_frames * i / num_frames) for i in range(num_frames)]
        
        for idx in frame_indices:
            cap.set(cv2.CAP_PROP_POS_FRAMES, idx)
            ret, frame = cap.read()
            if ret:
                frame_rgb = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
                img = Image.fromarray(frame_rgb)
                frames.append(img)
        
        cap.release()
        os.unlink(tfile.name)
        
        return frames
    except ImportError:
        st.error("opencv-python not installed. Run: pip install opencv-python")
        return []
    except Exception as e:
        st.error(f"Video processing error: {e}")
        return []

def analyze_video_with_gemini(video_bytes):
    try:
        from PIL import Image
        
        st.info("📹 Extracting frames from video...")
        frames = extract_video_frames(video_bytes, num_frames=8)
        
        if not frames:
            st.error("Failed to extract frames from video")
            return None
        
        st.success(f"✅ Extracted {len(frames)} frames from video")
        
        model = genai.GenerativeModel(
            model_name="gemini-2.0-flash",
            generation_config={
                "temperature": 0.7,
                "max_output_tokens": 4000,
            }
        )
        
        prompt = """Analyze this process video and extract detailed information for creating a Business Requirements Document (BRD).

Provide comprehensive analysis including:

1. **Process Overview**: What is the overall process shown?
2. **Step-by-Step Workflow**: List each step visible
3. **Systems/Applications**: Identify software, systems, or tools
4. **User Actions**: What actions does the user perform?
5. **Data Inputs**: What data is being entered?
6. **Outputs/Results**: What are the outcomes?
7. **Pain Points**: Inefficiencies or manual steps
8. **Business Rules**: Validation rules or business logic
9. **Integration Points**: Systems that communicate
10. **Automation Opportunities**: Which steps can be automated?

Be detailed and specific with clear formatting."""

        st.info("🤖 Analyzing video with Gemini AI...")
        response = model.generate_content([prompt] + frames)
        
        return response.text
        
    except Exception as e:
        st.error(f"Video analysis error: {e}")
        st.info("Tip: Ensure you're using gemini-1.5-flash or gemini-1.5-pro for video analysis")
        return None

# ---------------- FILE PROCESSING ----------------
def parse_docx_sections(path: str):
    doc = Document(path)
    sections = []
    current_title = "Introduction"
    current_paragraphs = []
    
    for p in doc.paragraphs:
        txt = p.text.strip()
        if not txt:
            continue
        
        style = (p.style.name or "").lower() if p.style else ""
        
        if style.startswith("heading") or (txt.isupper() and len(txt.split()) <= 6):
            if current_paragraphs:
                sections.append({
                    "title": current_title, 
                    "text": "\n\n".join(current_paragraphs)
                })
            current_title = txt
            current_paragraphs = []
        else:
            current_paragraphs.append(txt)
    
    if current_paragraphs:
        sections.append({
            "title": current_title, 
            "text": "\n\n".join(current_paragraphs)
        })
    
    return sections

def parse_excel_to_text(file_bytes: bytes, filename: str) -> str:
    try:
        df = pd.read_excel(io.BytesIO(file_bytes), sheet_name=None)
        text_parts = []
        
        for sheet_name, sheet_df in df.items():
            text_parts.append(f"\n## Sheet: {sheet_name}\n")
            
            cols = sheet_df.columns.tolist()
            has_question_col = any('q' in str(c).lower() or 'question' in str(c).lower() for c in cols)
            has_answer_col = any('a' in str(c).lower() or 'answer' in str(c).lower() for c in cols)
            
            if has_question_col and has_answer_col:
                qcol = next((c for c in cols if 'q' in str(c).lower() or 'question' in str(c).lower()), cols[0])
                acol = next((c for c in cols if 'a' in str(c).lower() or 'answer' in str(c).lower()), cols[1] if len(cols) > 1 else cols[0])
                
                for idx, row in sheet_df.iterrows():
                    q = str(row[qcol]) if pd.notna(row[qcol]) else ""
                    a = str(row[acol]) if pd.notna(row[acol]) else ""
                    if q or a:
                        text_parts.append(f"Q: {q}\nA: {a}\n")
            else:
                text_parts.append(sheet_df.to_string(index=False))
        
        return "\n".join(text_parts)
    except Exception as e:
        st.warning(f"Excel parsing warning: {e}")
        return ""

def parse_csv_to_text(file_bytes: bytes) -> str:
    try:
        df = pd.read_csv(io.BytesIO(file_bytes))
        cols = df.columns.tolist()
        has_question_col = any('q' in str(c).lower() or 'question' in str(c).lower() for c in cols)
        has_answer_col = any('a' in str(c).lower() or 'answer' in str(c).lower() for c in cols)
        
        if has_question_col and has_answer_col:
            qcol = next((c for c in cols if 'q' in str(c).lower() or 'question' in str(c).lower()), cols[0])
            acol = next((c for c in cols if 'a' in str(c).lower() or 'answer' in str(c).lower()), cols[1] if len(cols) > 1 else cols[0])
            
            parts = []
            for _, row in df.iterrows():
                q = str(row[qcol]) if pd.notna(row[qcol]) else ""
                a = str(row[acol]) if pd.notna(row[acol]) else ""
                if q or a:
                    parts.append(f"Q: {q}\nA: {a}\n")
            return "\n".join(parts)
        else:
            return df.to_string(index=False)
    except Exception as e:
        st.warning(f"CSV parsing warning: {e}")
        return ""

def parse_text_file(file_bytes: bytes) -> str:
    try:
        return file_bytes.decode('utf-8', errors='ignore')
    except Exception as e:
        st.warning(f"Text parsing warning: {e}")
        return ""

def parse_uploaded_file(uploaded_file) -> str:
    if uploaded_file is None:
        return ""
    
    file_bytes = uploaded_file.read()
    filename = uploaded_file.name
    ext = filename.split('.')[-1].lower()
    
    if ext in ['xlsx', 'xls']:
        return parse_excel_to_text(file_bytes, filename)
    elif ext == 'csv':
        return parse_csv_to_text(file_bytes)
    elif ext == 'txt':
        return parse_text_file(file_bytes)
    elif ext == 'docx':
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        tmp.write(file_bytes)
        tmp.close()
        try:
            sections = parse_docx_sections(tmp.name)
            return "\n\n".join([f"## {s['title']}\n{s['text']}" for s in sections])
        finally:
            os.unlink(tmp.name)
    else:
        st.warning(f"Unsupported file type: {ext}")
        return ""

def analyze_sample_brd(file_bytes):
    try:
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        tmp.write(file_bytes)
        tmp.close()
        
        doc = Document(tmp.name)
        
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        
        table_data = []
        for table in doc.tables:
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(" | ".join(row_data))
        
        os.unlink(tmp.name)
        
        combined_text = "\n".join(full_text)
        if table_data:
            combined_text += "\n\nTABLES:\n" + "\n".join(table_data)
        
        return combined_text
        
    except Exception as e:
        st.error(f"Sample file analysis error: {e}")
        return None

# ---------------- EMBEDDINGS ----------------
def embed_text(text: str):
    if embed_model is None:
        return None
    try:
        vec = embed_model.encode(text, show_progress_bar=False)
        return vec.tolist()
    except Exception as e:
        st.error(f"Embedding error: {e}")
        return None

def ingest_docx_bytes_to_store(file_bytes: bytes, filename: str):
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    tmp.write(file_bytes)
    tmp.flush()
    tmp.close()
    
    try:
        sections = parse_docx_sections(tmp.name)
        vectors = []
        
        for i, sec in enumerate(sections):
            sid = f"{Path(filename).stem}__{i}"
            emb = embed_text(sec["text"])
            
            if emb:
                metadata = {
                    "source": filename, 
                    "section_title": sec["title"], 
                    "text": sec["text"][:3000]
                }
                vectors.append((sid, emb, metadata))
        
        if vectors:
            vector_store.upsert(vectors)
            st.session_state.vector_store['vectors'] = vector_store.vectors
            return len(vectors)
        return 0
    finally:
        os.unlink(tmp.name)

def retrieve_similar_sections(query: str, top_k: int = 6):
    q_emb = embed_text(query)
    if not q_emb:
        return []
    
    resp = vector_store.query(vector=q_emb, top_k=top_k)
    results = []
    
    for m in resp["matches"]:
        meta = m.get("metadata", {})
        results.append({
            "id": m.get("id"),
            "score": m.get("score"),
            "source": meta.get("source"),
            "section_title": meta.get("section_title"),
            "text": meta.get("text", "")
        })
    
    return results

# ---------------- GEMINI API ----------------
def call_gemini(prompt: str, max_tokens: int = 4000):
    try:
        model = genai.GenerativeModel(
            model_name=GEMINI_MODEL,
            generation_config={
                "temperature": 0.7,
                "top_p": 0.95,
                "top_k": 40,
                "max_output_tokens": max_tokens,
            }
        )
        
        response = model.generate_content(prompt)
        return response.text
        
    except Exception as e:
        error_msg = str(e)
        st.error(f"❌ Gemini API Error: {error_msg}")
        
        if "model name" in error_msg.lower():
            st.info("""
            **Try these models:**
            - gemini-2.5-flash
        
            """)
        
      
      
        return None
# ---------------- FLOW DIAGRAM ----------------
import requests
def generate_mermaid_flowchart(process_flow_steps):
    prompt = f"""
You are an expert business analyst. Given these process flow steps, generate a flowchart in Mermaid format that visually represents the workflow.
Use appropriate shapes for start/end, decisions, and actions. Return ONLY the Mermaid code, no explanations.

Process Flow Steps:
{chr(10).join([f"{i+1}. {step}" for i, step in enumerate(process_flow_steps)])}
"""
    mermaid_code = call_gemini(prompt, max_tokens=1000)
    return mermaid_code

def mermaid_to_image(mermaid_code: str, output_img_path: str) -> Optional[str]:
    """
    Renders Mermaid code to PNG using mermaid-py locally.
    Suppresses detailed errors on frontend.
    """
    try:
        # Validate Mermaid syntax
        if not mermaid_code.strip().startswith(('graph TD', 'graph LR', 'graph TB')):
            # Instead of raising ValueError, log internally and return None
            return None

        mermaid = Mermaid(mermaid_code)
        svg_data = mermaid.to_svg()
        if not svg_data or not isinstance(svg_data, str):
            return None
        
        png_data = BytesIO()
        svg2png(bytestring=svg_data.encode('utf-8'), write_to=png_data)
        png_data.seek(0)
        with open(output_img_path, "wb") as f:
            f.write(png_data.getvalue())
        return output_img_path
    except Exception as e:
        # Suppress specific error details on frontend
        st.warning("An error occurred while generating the flowchart. Please try again.")
        return None


def add_page_border(section):
    # Add border to the section using XML
    border_xml = r'''
    <w:pgBorders %s>
        <w:top w:val="single" w:sz="18" w:space="24" w:color="000000"/>
        <w:left w:val="single" w:sz="18" w:space="24" w:color="000000"/>
        <w:bottom w:val="single" w:sz="18" w:space="24" w:color="000000"/>
        <w:right w:val="single" w:sz="18" w:space="24" w:color="000000"/>
    </w:pgBorders>
    ''' % nsdecls('w')
    sectPr = section._sectPr
    sectPr.append(parse_xml(border_xml))
    
DEFAULT_HEADER_IMAGE_PATH = os.path.join("assets", "ess_logo.png")




# ---------------- BRD GENERATION ----------------
def create_formatted_brd(brd_data: dict, output_path: str,header_image_path=DEFAULT_HEADER_IMAGE_PATH):
    doc = Document()

    section = doc.sections[0]
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    add_page_border(section)

    if header_image_path and os.path.exists(header_image_path):
        header = section.header
        header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        run = header_para.add_run()
        try:
            run.add_picture(header_image_path, width=Inches(1.2))
            header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        except Exception as e:
            header_para.add_run().add_text(f"Error adding logo: {str(e)}")
    else:
        if header_image_path:
            header = section.header
            header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            header_para.add_run().add_text("Logo not found at specified path.")
    
    

 
    
     

    
    
  
    

   
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'times new roman'
    font.size = Pt(13)
    
    title = doc.add_heading('BRD (Business Requirement Document)', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f"• COMPANY NAME: {brd_data.get('company_name', 'Not Specified')}")
    doc.add_paragraph(f"• SOW VERSION: {brd_data.get('sow_version', 'V1.0')}")
    doc.add_paragraph(f"• DATE SUBMITTED: {brd_data.get('date_submitted', datetime.now().strftime('%d %B %Y'))}")
    doc.add_paragraph()
    doc.add_paragraph("• ORGANIZATION")
    doc.add_paragraph()
    doc.add_paragraph(f"• CLIENT NAME: {brd_data.get('client_name', 'Not Specified')}")
    doc.add_paragraph()
    doc.add_paragraph("• PROJECT")
    doc.add_paragraph()
    doc.add_paragraph(f"• PROJECT NAME: {brd_data.get('project_name', 'Not Specified')}")
    
    doc.add_page_break()
    
    doc.add_heading('Contents', level=1)
    toc_items = ['Overview', 'Purpose', 'Business Objectives', 'Current Challenges',
                 'Introduction', 'Scope', 'Process Description', 'Stakeholders',
                 'Process Flow', 'Functional Requirements', 'Data Validation',
                 'Error Handling', 'Required Inputs', 'Detailed Steps', 'Output',
                 'Result', 'Non-Functional Requirements', 'Dependencies',
                 'Assumptions', 'Document History']
    for item in toc_items:
        doc.add_paragraph(f"• {item}", style='List Bullet')
    
    doc.add_page_break()
    
    doc.add_heading('Overview', level=1)
    doc.add_paragraph(brd_data.get('overview', 'Process overview.'))
    
    doc.add_heading('Purpose', level=1)
    doc.add_paragraph(brd_data.get('purpose', 'Document purpose.'))
    
    doc.add_heading('Business Objectives', level=1)
    for obj in brd_data.get('business_objectives', []):
        doc.add_paragraph(obj, style='List Bullet')
    
    doc.add_heading('Current Challenges', level=1)
    for challenge in brd_data.get('current_challenges', []):
        doc.add_paragraph(challenge, style='List Bullet')
    
    doc.add_page_break()
    
    doc.add_heading('Introduction', level=1)
    doc.add_paragraph(brd_data.get('introduction', 'Introduction.'))
    
    doc.add_heading('Scope', level=1)
    for i, item in enumerate(brd_data.get('scope', []), 1):
        if isinstance(item, dict):
            p = doc.add_paragraph(f"{i}. ")
            p.add_run(item.get('title', '')).bold = True
            for detail in item.get('details', []):
                doc.add_paragraph(detail, style='List Bullet')
        else:
            doc.add_paragraph(f"{i}. {item}")
    
    doc.add_heading('Process Description', level=1)
    doc.add_paragraph(brd_data.get('process_description', ''))
    
    doc.add_heading('Stakeholders', level=2)
    for i, stakeholder in enumerate(brd_data.get('stakeholders', []), 1):
        if isinstance(stakeholder, dict):
            doc.add_paragraph(f"{i}. {stakeholder.get('role', '')}: {stakeholder.get('description', '')}")
        else:
            doc.add_paragraph(f"{i}. {stakeholder}")
    
    doc.add_heading('Process Flow', level=1)
    for step in brd_data.get('process_flow', []):
        doc.add_paragraph(step, style='List Number')
    
    doc.add_page_break()
    
    doc.add_heading('Functional Requirements', level=1)
    for req in brd_data.get('functional_requirements', []):
        p = doc.add_paragraph()
        p.add_run(req).bold = True
    
    doc.add_heading('Data Validation', level=1)
    for i, validation in enumerate(brd_data.get('data_validations', []), 1):
        if isinstance(validation, dict):
            p = doc.add_paragraph(f"{i}. ")
            p.add_run(validation.get('rule', '')).bold = True
            for detail in validation.get('details', []):
                doc.add_paragraph(detail, style='List Bullet')
        else:
            doc.add_paragraph(f"{i}. {validation}")
    
    doc.add_heading('Error Handling', level=2)
    for i, error in enumerate(brd_data.get('error_handling', []), 1):
        doc.add_paragraph(f"{i}. {error}")
    
    doc.add_heading('Required Inputs', level=1)
    for inp in brd_data.get('required_inputs', []):
        doc.add_paragraph(inp, style='List Number')
    
    doc.add_heading('Detailed Process Steps', level=1)
    for step in brd_data.get('detailed_process_steps', []):
        if isinstance(step, dict):
            doc.add_paragraph(f"{step.get('step_number', '')}. {step.get('description', '')}")
            for detail in step.get('details', []):
                doc.add_paragraph(detail, style='List Bullet')
            if 'note' in step:
                p = doc.add_paragraph()
                p.add_run('NOTE: ').bold = True
                p.add_run(step['note'])
        else:
            doc.add_paragraph(str(step))
    
    doc.add_page_break()
    
    doc.add_heading('Output', level=1)
    for output in brd_data.get('outputs', []):
        doc.add_paragraph(output, style='List Bullet')
    
    doc.add_heading('Result', level=1)
    for result in brd_data.get('results', []):
        doc.add_paragraph(result, style='List Bullet')
    
    doc.add_heading('Non-Functional Requirements', level=1)
    for category, items in brd_data.get('non_functional_requirements', {}).items():
        doc.add_heading(category, level=2)
        for item in items:
            doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('Dependencies', level=1)
    doc.add_heading('External', level=2)
    for dep in brd_data.get('external_dependencies', []):
        doc.add_paragraph(dep, style='List Bullet')
    
    doc.add_heading('Internal', level=2)
    for dep in brd_data.get('internal_dependencies', []):
        doc.add_paragraph(dep, style='List Bullet')
    
    doc.add_heading('Assumptions', level=1)
    for assumption in brd_data.get('assumptions_constraints', []):
        doc.add_paragraph(assumption, style='List Bullet')
    
    doc.add_heading('Document History', level=1)
    table = doc.add_table(rows=2, cols=6)
    table.style = 'Light Grid Accent 1'
    
    hdr_cells = table.rows[0].cells
    headers = ['Date', 'Version', 'Prepared By', 'Reviewed By', 'Approved By', 'Business User']
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
    
    row_cells = table.rows[1].cells
    history = brd_data.get('document_history', {})
    row_cells[0].text = history.get('date', datetime.now().strftime('%d.%m.%Y'))
    row_cells[1].text = history.get('version', 'V1.0')
    row_cells[2].text = history.get('prepared_by', '')
    row_cells[3].text = history.get('reviewed_by', '')
    row_cells[4].text = history.get('approved_by', '')
    row_cells[5].text = history.get('business_user', '')

    process_flow_steps = brd_data.get('process_flow', [])
    if process_flow_steps:
        mermaid_code = generate_mermaid_flowchart(process_flow_steps)
        if mermaid_code:
            diagram_path = os.path.join(tempfile.gettempdir(), "process_flow_mermaid.png")
            img_path = mermaid_to_image(mermaid_code, diagram_path)
            if img_path and os.path.exists(img_path):
                doc.add_page_break()
                doc.add_heading('Process Flow Diagram', level=1)
                doc.add_paragraph("Below is the process flow diagram generated by Gemini AI.")
                doc.add_picture(img_path, width=Inches(6))
            else:
                doc.add_page_break()
                doc.add_heading('Process Flow Diagram (Mermaid)', level=1)
                doc.add_paragraph("Could not render diagram image. Here is the Mermaid code:")
                para = doc.add_paragraph()
                run = para.add_run(mermaid_code)
                run.font.name = 'Consolas'
                run.font.size = Pt(10)
    
    doc.save(output_path)
    return output_path

# ---------------- UI ----------------
st.header("📚 Step 1: Knowledge Base (Optional)")

with st.expander("About Knowledge Base", expanded=False):
    st.info("Upload existing BRDs to teach AI your standards.")

col1, col2 = st.columns([2, 1])

with col1:
    brd_uploads = st.file_uploader(
        "Upload BRDs (.docx)",
        type=["docx"],
        accept_multiple_files=True,
        key="brd_ingest"
    )

with col2:
    st.metric("Documents", len(vector_store.vectors))
    if st.button("Clear KB", use_container_width=True):
        vector_store.clear()
        st.session_state.vector_store['vectors'] = []
        st.success("Cleared!")
        st.rerun()

if brd_uploads and st.button("Ingest", type="primary", use_container_width=True):
    total = 0
    progress = st.progress(0)
    
    for i, f in enumerate(brd_uploads):
        try:
            content = f.read()
            cnt = ingest_docx_bytes_to_store(content, f.name)
            total += cnt
            st.success(f"✅ {cnt} sections from {f.name}")
        except Exception as e:
            st.error(f"❌ Failed: {f.name}: {e}")
        
        progress.progress((i + 1) / len(brd_uploads))
    
    progress.empty()
    st.success(f"Ingested {total} sections!")

st.markdown("---")

st.header("🎯 Step 2: Generate BRD")

generation_mode = st.radio(
    "Choose Mode:",
    options=["Create from Scratch", "Use Sample Template", "Analyze Video"],
    horizontal=True,
    key="generation_mode"
)

st.markdown("---")

# SAMPLE TEMPLATE MODE - OUTSIDE FORM
if generation_mode == "Use Sample Template":
    st.subheader("📋 Sample Template Mode")
    st.info("Upload a sample BRD (.docx) to use as template. The AI will analyze its structure and generate a new BRD following the same format.")
    
    sample_file = st.file_uploader(
        "Upload Sample BRD Template (.docx)",
        type=["docx"],
        key="sample_brd_upload"
    )
    
    if sample_file:
        if st.button("📖 Analyze Sample Template", type="primary"):
            with st.spinner("Analyzing sample BRD structure..."):
                file_bytes = sample_file.read()
                sample_content = analyze_sample_brd(file_bytes)
                
                if sample_content:
                    st.session_state.sample_content = sample_content
                    st.success(f"✅ Analyzed {len(sample_content)} characters from template")
                    with st.expander("👁️ Preview Sample Content"):
                        st.text_area("Template Content", sample_content[:3000] + "..." if len(sample_content) > 3000 else sample_content, height=300, disabled=True)
                else:
                    st.error("Failed to analyze sample file")
    
    if st.session_state.sample_content:
        st.success("✅ Sample template is ready! Fill in project details below and generate BRD.")

# VIDEO ANALYSIS MODE - OUTSIDE FORM
elif generation_mode == "Analyze Video":
    st.subheader("🎥 Video Analysis Mode")
    st.info("Upload a process video (MP4, AVI, MOV, MKV). AI will analyze the video and generate BRD based on visual workflow.")
    st.warning("⚠️ Requires: pip install opencv-python")
    
    video_file = st.file_uploader(
        "Upload Process Video",
        type=["mp4", "avi", "mov", "mkv"],
        key="video_upload"
    )
    
    if video_file:
        col1, col2 = st.columns([2, 1])
        
        with col1:
            st.video(video_file)
        
        with col2:
            video_bytes = video_file.getvalue()
            st.metric("Video Size", f"{len(video_bytes) / (1024*1024):.1f} MB")
            
            if st.button("🎬 Analyze Video Now", type="primary"):
                with st.spinner("Processing video... This may take 1-2 minutes"):
                    video_analysis = analyze_video_with_gemini(video_bytes)
                    
                    if video_analysis:
                        st.session_state.video_analysis = video_analysis
                        st.success("✅ Video analysis completed!")
                        with st.expander("📊 Analysis Results", expanded=True):
                            st.markdown(video_analysis)
                    else:
                        st.error("Failed to analyze video")
    
    if st.session_state.video_analysis:
        st.success("✅ Video analysis is ready! Fill in project details below and generate BRD.")

# PROCESS FILE UPLOAD - OUTSIDE FORM (for Create from Scratch mode)
elif generation_mode == "Create from Scratch":
    st.subheader("📄 Process Details Input")
    
    input_method = st.radio(
        "How do you want to provide process details?",
        options=["Type Manually", "Upload File"],
        horizontal=True,
        key="input_method"
    )
    
    if input_method == "Upload File":
        st.info("Upload a file containing process details (Excel Q&A, Word doc, CSV, or Text file)")
        
        process_file = st.file_uploader(
            "Upload Process Details File",
            type=["xlsx", "xls", "csv", "txt", "docx"],
            key="process_details_file"
        )
        
        if process_file:
            if st.button("📖 Read Process File", type="primary"):
                with st.spinner("Reading and parsing file..."):
                    parsed_content = parse_uploaded_file(process_file)
                    
                    if parsed_content and len(parsed_content) > 50:
                        st.session_state.process_file_content = parsed_content
                        st.success(f"✅ Read {len(parsed_content)} characters from {process_file.name}")
                        with st.expander("👁️ Preview Extracted Content"):
                            st.text_area("File Content", parsed_content[:3000] + "..." if len(parsed_content) > 3000 else parsed_content, height=300, disabled=True)
                    else:
                        st.error("Could not extract sufficient content from file. Please check file format.")
        
        if st.session_state.process_file_content:
            st.success("✅ Process file is ready! Fill in project details below and generate BRD.")

st.markdown("---")

# MAIN FORM - PROJECT DETAILS
with st.form("brd_generation_form"):
    st.subheader("📋 Project Information")
    
    col1, col2, col3 = st.columns(3)
    
    with col1:
        company_name = st.text_input("Company Name", value="Eastern Software Solutions Pvt. Ltd.")
        client_name = st.text_input("Client Name")
    
    with col2:
        project_name = st.text_input("Project Name *", placeholder="e.g., Inventory Replenishment")
        sow_version = st.text_input("Version", value="V1.0")
    
    with col3:
        prepared_by = st.text_input("Prepared By")
        reviewed_by = st.text_input("Reviewed By")
    
    st.markdown("---")
    
    # Process description input - only show text area if typing manually
    if generation_mode == "Create from Scratch":
        if 'input_method' in st.session_state and st.session_state.get('input_method') == "Type Manually":
            st.subheader("✍️ Type Process Description")
            process_description = st.text_area(
                "Process Description *",
                height=250,
                placeholder="Describe the process to be automated in detail...\n\nExample:\n- What is the current process?\n- What are the pain points?\n- What should be automated?\n- What are the expected outcomes?",
                help="Provide comprehensive process details (minimum 50 characters)"
            )
        else:
            process_description = ""
            if st.session_state.process_file_content:
                st.info(f"📄 Using uploaded file content ({len(st.session_state.process_file_content)} characters)")
            else:
                st.warning("⚠️ Please upload a process file above before generating")
    else:
        st.subheader("➕ Additional Context (Optional)")
        process_description = st.text_area(
            "Additional Details",
            height=100,
            placeholder="Add any additional context or specific requirements...",
            help="Optional additional information to supplement the template or video analysis"
        )
    
    st.markdown("---")
    st.subheader("📎 Additional Reference Files (Optional)")
    
    additional_qa_file = st.file_uploader(
        "Upload Additional Q&A or Reference Files",
        type=["xlsx", "xls", "csv", "txt", "docx"],
        help="Optional: Upload additional reference materials",
        key="additional_qa_upload"
    )
    
    col1, col2 = st.columns([3, 1])
    with col1:
        generate_button = st.form_submit_button(
            "🚀 Generate BRD Document",
            type="primary",
            use_container_width=True
        )
    with col2:
        st.caption("* Required fields")

# HANDLE BRD GENERATION
if generate_button:
    # Validation
    validation_passed = True
    
    if not project_name:
        st.error("⚠️ Please provide a Project Name")
        validation_passed = False
    
    elif generation_mode == "Create from Scratch":
        if 'input_method' in st.session_state and st.session_state.get('input_method') == "Type Manually":
            if not process_description or len(process_description.strip()) < 50:
                st.error("⚠️ Please provide process details (minimum 50 characters)")
                validation_passed = False
        else:
            if not st.session_state.process_file_content:
                st.error("⚠️ Please upload and read a process file first")
                validation_passed = False
    
    elif generation_mode == "Use Sample Template":
        if not st.session_state.sample_content:
            st.error("⚠️ Please upload and analyze a sample BRD template first")
            validation_passed = False
    
    elif generation_mode == "Analyze Video":
        if not st.session_state.video_analysis:
            st.error("⚠️ Please upload and analyze a video first")
            validation_passed = False
    
    if validation_passed:
        with st.spinner("🤖 AI is generating your comprehensive BRD... This may take 30-90 seconds"):
            
            # Prepare process description based on mode
            final_process_description = ""
            
            if generation_mode == "Create from Scratch":
                if st.session_state.process_file_content:
                    final_process_description = st.session_state.process_file_content
                else:
                    final_process_description = process_description
            elif generation_mode == "Use Sample Template":
                final_process_description = process_description if process_description else "Use the sample template structure"
            elif generation_mode == "Analyze Video":
                final_process_description = st.session_state.video_analysis
                if process_description:
                    final_process_description = f"{process_description}\n\nVIDEO ANALYSIS:\n{st.session_state.video_analysis}"
            
            # Build context
            context_parts = []
            
            # Add knowledge base context
            if vector_store.vectors and generation_mode != "Use Sample Template":
                with st.spinner("🔍 Searching knowledge base..."):
                    similar = retrieve_similar_sections(final_process_description, top_k=6)
                    if similar:
                        context_parts.append("## Similar BRD Sections from Knowledge Base:\n")
                        for s in similar:
                            context_parts.append(f"### {s['section_title']} (from {s['source']})\n{s['text']}\n")
                        st.info(f"Found {len(similar)} relevant sections from knowledge base")
            
            # Add additional Q&A file
            if additional_qa_file:
                with st.spinner("📖 Processing additional reference file..."):
                    qa_content = parse_uploaded_file(additional_qa_file)
                    if qa_content:
                        context_parts.append(f"\n## Additional Reference Material:\n{qa_content}\n")
                        st.info(f"Added {len(qa_content)} characters from {additional_qa_file.name}")
            
            context = "\n".join(context_parts) if context_parts else ""
            
            # Generate prompt based on mode
            if generation_mode == "Use Sample Template":
                prompt = f"""You are an expert Business Analyst. Create a NEW BRD based on the SAMPLE template structure.

SAMPLE BRD TEMPLATE (follow this exact structure):
{st.session_state.sample_content[:15000]}

NEW PROJECT INFORMATION:
- Company: {company_name}
- Client: {client_name}
- Project Name: {project_name}
- Version: {sow_version}
- Process Description: {final_process_description}

INSTRUCTIONS:
1. Analyze the SAMPLE BRD structure carefully
2. Maintain the EXACT same section organization
3. Use similar writing style and tone
4. Fill in content specific to the NEW project
5. Keep the same level of detail and professionalism

Return a complete BRD in JSON format with all standard sections. Use this structure:
{{
  "overview": "Write overview based on new project",
  "purpose": "Purpose statement",
  "business_objectives": ["Objective 1", "Objective 2", "Objective 3", "Objective 4", "Objective 5"],
  "current_challenges": ["Challenge 1", "Challenge 2", "Challenge 3", "Challenge 4"],
  "introduction": "Introduction text",
  "scope": [{{"title": "Scope item", "details": ["Detail 1", "Detail 2"]}}, ...],
  "process_description": "Process description",
  "stakeholders": [{{"role": "Role name", "description": "Description"}}, ...],
  "process_flow": ["Step 1", "Step 2", "Step 3", ...],
  "functional_requirements": ["Requirement 1", "Requirement 2", ...],
  "data_validations": [{{"rule": "Rule name", "details": ["Detail 1", "Detail 2"]}}, ...],
  "error_handling": ["Error handling 1", "Error handling 2", ...],
  "required_inputs": ["Input 1", "Input 2", ...],
  "detailed_process_steps": [{{
    "step_number": 1,
    "description": "Step description",
    "details": ["Detail 1", "Detail 2"],
    "note": "Optional note"
  }}, ...],
  "outputs": ["Output 1", "Output 2", ...],
  "results": ["Result 1", "Result 2", ...],
  "non_functional_requirements": {{
    "Security": ["Security req 1", "Security req 2"],
    "Performance": ["Performance req 1", "Performance req 2"],
    "Reliability": ["Reliability req 1", "Reliability req 2"]
  }},
  "external_dependencies": ["Dependency 1", "Dependency 2", ...],
  "internal_dependencies": ["Dependency 1", "Dependency 2", ...],
  "assumptions_constraints": ["Assumption 1", "Assumption 2", ...]
}}

Return ONLY valid JSON without markdown formatting or code blocks."""

            else:
                # Standard prompt for other modes
                prompt = f"""You are an expert Business Analyst with 15+ years of experience creating comprehensive BRDs for automation projects.

PROJECT INFORMATION:
- Company: {company_name}
- Client: {client_name}
- Project Name: {project_name}
- Version: {sow_version}

PROCESS DETAILS:
{final_process_description}

REFERENCE CONTEXT:
{context}

Create a detailed, professional BRD in JSON format. Include ALL standard sections with substantial, specific content based on the process details provided.

Return JSON with this exact structure:
{{
  "overview": "3-4 paragraph comprehensive overview of the project and automation strategy",
  "purpose": "2-3 paragraph purpose statement explaining document objectives",
  "business_objectives": [
    "Business objective 1 - specific and measurable",
    "Business objective 2 - focus on business value",
    "Business objective 3 - efficiency gains",
    "Business objective 4 - risk reduction",
    "Business objective 5 - scalability"
  ],
  "current_challenges": [
    "Challenge 1: Detailed pain point description",
    "Challenge 2: Another business challenge",
    "Challenge 3: Inefficiency details",
    "Challenge 4: Error-prone area"
  ],
  "introduction": "2-3 paragraph introduction setting context",
  "scope": [
    {{"title": "Real-time monitoring and tracking", "details": ["Detail about monitoring", "What data is tracked", "Benefits"]}},
    {{"title": "Automated processing and validation", "details": ["What gets automated", "How validation works", "Error detection"]}},
    {{"title": "Exception handling and notifications", "details": ["How exceptions detected", "Notification mechanisms", "Escalation procedures"]}}
  ],
  "process_description": "3-4 paragraph detailed end-to-end workflow description",
  "stakeholders": [
    {{"role": "Process Owner/Team Leader", "description": "Detailed responsibilities and involvement"}},
    {{"role": "IT/RPA Development Team", "description": "Their support role"}},
    {{"role": "Operations Team", "description": "Day-to-day operations role"}}
  ],
  "process_flow": [
    "Step 1: Detailed description of first step",
    "Step 2: What happens next in workflow",
    "Step 3: Continue logical flow",
    "Step 4: Include decision points",
    "Step 5: Describe final steps"
  ],
  "functional_requirements": [
    "Automatically validate data against business rules with real-time integration",
    "Generate and submit processing requests with full audit trail",
    "Consolidate multiple requests into optimized batches",
    "Maintain comprehensive logs with timestamps",
    "Send automated exception alerts to stakeholders",
    "Provide detailed daily reports with metrics"
  ],
  "data_validations": [
    {{"rule": "Verify data completeness and accuracy", "details": ["Check required fields", "Validate data formats", "Missing data triggers exception"]}},
    {{"rule": "Validate system mappings and configurations", "details": ["Ensure codes exist in master data", "Verify mappings", "Invalid mappings generate errors"]}},
    {{"rule": "Confirm data consistency and integrity", "details": ["Cross-check related data", "Validate business rules", "Flag inconsistencies"]}}
  ],
  "error_handling": [
    "System/Application Error: Retry 3 times at 5-minute intervals before failing and notifying support",
    "Data Validation Error: Log with details and notify stakeholders immediately",
    "Business Rule Violation: Capture in exception log and escalate by severity"
  ],
  "required_inputs": [
    "System credentials with appropriate access levels",
    "Current data from source systems or files",
    "Configuration parameters and business rules",
    "Notification distribution lists"
  ],
  "detailed_process_steps": [
    {{
      "step_number": 1,
      "description": "BOT logs into system using provided credentials",
      "details": ["Navigate to login page", "Enter credentials securely", "Validate successful login"]
    }},
    {{
      "step_number": 2,
      "description": "BOT navigates to required module/section",
      "details": ["Locate appropriate menu", "Wait for page load", "Verify correct page"]
    }},
    {{
      "step_number": 3,
      "description": "BOT extracts and validates data",
      "details": ["Read data from source", "Perform validation checks", "Identify issues"],
      "note": "Validation failures are logged and stakeholders notified"
    }},
    {{
      "step_number": 4,
      "description": "BOT processes validated records",
      "details": ["Apply business logic", "Perform calculations", "Prepare data"]
    }},
    {{
      "step_number": 5,
      "description": "BOT submits data to target system",
      "details": ["Enter data in fields", "Validate entries", "Submit and capture confirmation"]
    }},
    {{
      "step_number": 6,
      "description": "BOT generates reports and notifications",
      "details": ["Compile statistics", "Create summary report", "Send emails", "Log completion"]
    }}
  ],
  "outputs": [
    "Successfully processed transactions in target system",
    "Comprehensive processing report with statistics",
    "Exception log with failure details",
    "Email notifications with summary"
  ],
  "results": [
    "Faster processing time - reduction from hours to minutes",
    "Improved accuracy - elimination of manual errors",
    "Enhanced visibility - real-time tracking",
    "Better compliance - complete audit trail",
    "Reduced operational costs - automation of repetitive tasks"
  ],
  "non_functional_requirements": {{
    "Security": [
      "All credentials stored securely using encryption",
      "Role-based access controls",
      "Audit logs for all system actions",
      "Compliance with data protection regulations"
    ],
    "Performance": [
      "Process completion within SLA timeframes",
      "Handle expected transaction volumes",
      "Scalable to support business growth",
      "Minimal impact on system resources"
    ],
    "Reliability": [
      "99% uptime target for automated processes",
      "Robust error handling and recovery",
      "Automated monitoring and alerting",
      "Scheduled execution with retry capabilities"
    ],
    "User-Interface": [
      "Intuitive dashboards for monitoring",
      "Clear exception reporting workflows",
      "Easy access to logs and audit trails"
    ]
  }},
  "external_dependencies": [
    "Integration with target system - requires stable API or UI access",
    "Email system for sending notifications and reports",
    "Source data availability on schedule",
    "Third-party services required for processing"
  ],
  "internal_dependencies": [
    "User accounts with appropriate permissions for all systems",
    "Network connectivity between automation server and target systems",
    "Proper configuration of business rules and parameters",
    "Availability of support team for exception handling"
  ],
  "assumptions_constraints": [
    "Source data is available in expected format and location at scheduled time",
    "System interfaces and APIs remain stable without breaking changes",
    "Users have valid credentials with necessary access rights",
    "Network connectivity is reliable during processing windows",
    "Support team is available during business hours for exception handling"
  ]
}}

CRITICAL: Return ONLY valid JSON, no markdown formatting, no code blocks, no explanations. Be specific to the actual process described."""

            # Call Gemini API
            with st.spinner("💭 Gemini AI is analyzing and generating..."):
                response = call_gemini(prompt, max_tokens=8000)
            
            if response:
                # Parse JSON response
                json_data = None
                
                try:
                    json_data = json.loads(response)
                    st.success("✅ Successfully parsed AI response")
                    
                except json.JSONDecodeError:
                    # Try to extract JSON from markdown code blocks
                    json_match = re.search(r'```(?:json)?\s*(\{.*?\})\s*```', response, re.DOTALL)
                    if json_match:
                        try:
                            json_data = json.loads(json_match.group(1))
                            st.success("✅ Successfully extracted JSON from response")
                        except:
                            st.warning("⚠️ Found JSON block but couldn't parse it")
                    
                    # Last resort: find any JSON object
                    if not json_data:
                        json_match = re.search(r'(\{[^{}]*(?:\{[^{}]*\}[^{}]*)*\})', response, re.DOTALL)
                        if json_match:
                            try:
                                json_data = json.loads(json_match.group(1))
                                st.success("✅ Successfully extracted JSON")
                            except:
                                pass
                
                if json_data:
                    # Add metadata
                    json_data['company_name'] = company_name
                    json_data['client_name'] = client_name
                    json_data['project_name'] = project_name
                    json_data['sow_version'] = sow_version
                    json_data['date_submitted'] = datetime.now().strftime('%d %B %Y')
                    json_data['document_history'] = {
                        'date': datetime.now().strftime('%d.%m.%Y'),
                        'version': sow_version,
                        'prepared_by': prepared_by,
                        'reviewed_by': reviewed_by,
                        'approved_by': '',
                        'business_user': ''
                    }
                    
                    # Generate Word document
                    try:
                        with st.spinner("📝 Creating Word document..."):
                            output_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx", mode='wb')
                            output_path = output_file.name
                            output_file.close()
                            
                            create_formatted_brd(json_data, output_path)
                            
                            with open(output_path, 'rb') as f:
                                doc_bytes = f.read()
                            
                            os.unlink(output_path)
                        
                        # Success!
                        st.balloons()
                        st.success("🎉 BRD Document Generated Successfully!")
                        
                        col1, col2 = st.columns([2, 1])
                        
                        with col1:
                            filename = f"BRD_{project_name.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
                            st.download_button(
                                label="📥 Download BRD Document",
                                data=doc_bytes,
                                file_name=filename,
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                use_container_width=True,
                                type="primary"
                            )
                        
                        with col2:
                            st.metric("Document Size", f"{len(doc_bytes) // 1024} KB")
                        
                        with st.expander("📄 Preview Generated Content Structure", expanded=True):
                            st.json(json_data)
                        
                        with st.expander("📊 Document Statistics"):
                            col1, col2, col3 = st.columns(3)
                            
                            with col1:
                                st.metric("Business Objectives", len(json_data.get('business_objectives', [])))
                                st.metric("Stakeholders", len(json_data.get('stakeholders', [])))
                            
                            with col2:
                                st.metric("Process Steps", len(json_data.get('detailed_process_steps', [])))
                                st.metric("Validations", len(json_data.get('data_validations', [])))
                            
                            with col3:
                                st.metric("Functional Requirements", len(json_data.get('functional_requirements', [])))
                                st.metric("Scope Items", len(json_data.get('scope', [])))
                        
                    except Exception as e:
                        st.error(f"❌ Error creating Word document: {e}")
                        st.info("The JSON data was generated successfully. You can view it below:")
                        st.json(json_data)
                        
                else:
                    st.error("❌ Could not parse AI response as JSON")
                    with st.expander("🔍 View Raw AI Response"):
                        st.code(response, language="text")
                    
                    st.info("""
                    **Troubleshooting Tips:**
                    1. Try simplifying your process description
                    2. Make sure your API key is valid
                    3. Check if you've exceeded API quota
                    4. Try again - sometimes the API returns inconsistent formats
                    """)
            else:
                st.error("❌ No response received from Gemini API")

st.markdown("---")

# Help Section
with st.expander("ℹ️ How to Use This Tool"):
    st.markdown("""
    ### Step-by-Step Guide:
    
    **Step 1: Build Knowledge Base (Optional)**
    - Upload existing BRD documents
    - AI learns your organization's style
    
    **Step 2: Choose Generation Mode**
    
    **🆕 Create from Scratch:**
    - Type process description manually OR
    - Upload process file (Excel, Word, CSV, Text)
    - Click "Read Process File" button to parse
    - Then generate BRD
    
    **📋 Use Sample Template:**
    - Upload your existing BRD template (.docx)
    - Click "Analyze Sample Template"
    - Fill in new project details
    - Generate BRD matching template structure
    
    **🎥 Analyze Video:**
    - Upload process recording (MP4, AVI, MOV, MKV)
    - Click "Analyze Video Now"
    - Fill in project details
    - Generate BRD from video analysis
    
    **Step 3: Fill Project Details**
    - Company, client, project name, etc.
    
    **Step 4: Generate & Download**
    - Click "Generate BRD Document"
    - Download Word document
    
    ### Supported File Formats:
    - Excel (.xlsx, .xls) - Q&A format
    - CSV - Data tables
    - Word (.docx) - Documentation or templates
    - Text (.txt) - Plain text descriptions
    - Video (.mp4, .avi, .mov, .mkv) - Process recordings
    
    ### Tips for Best Results:
    - **For File Uploads:** Click the parse/analyze button before generating
    - **For Videos:** Use 1080p resolution, 2-5 minutes duration
    - **For Templates:** Use recently approved, complete BRDs
    - **For Manual Input:** Be detailed and specific (50+ characters)
    """)

with st.expander("🔧 Configuration & Troubleshooting"):
    st.markdown(f"""
    ### Current Configuration:
    - **Model:** {GEMINI_MODEL}
    - **Embedding:** {EMBED_MODEL_NAME}
    - **API Key:** {"✅ Set" if GOOGLE_API_KEY else "❌ Not Set"}
    - **Knowledge Base:** {len(vector_store.vectors)} sections
    
    ### Supported Models:
    - gemini-2.0-flash-exp (latest)
    - gemini-1.5-flash (stable, supports vision)
    - gemini-1.5-pro (most capable)
    
    ### Setup .env File:
    ```
    GOOGLE_API_KEY=your_api_key_here
    GEMINI_MODEL=gemini-2.0-flash-exp
    EMBED_MODEL_NAME=all-MiniLM-L6-v2
    ```
    
    ### Install Dependencies:
    ```bash
    pip install streamlit python-dotenv python-docx pandas
    pip install openpyxl sentence-transformers google-generativeai
    pip install opencv-python  # For video analysis
    ```
    
    ### Common Issues:
    - **File upload not working:** Make sure to click the "Read" or "Analyze" button
    - **Video analysis fails:** Install opencv-python
    - **Model error:** Use gemini-1.5-flash for video analysis
    - **API quota exceeded:** Wait a few minutes and try again
    """)

st.markdown("---")
st.markdown("""
<div style='text-align: center; color: #666; padding: 20px;'>
    <p><strong>BRD Automation Studio V1.0</strong> | Powered by Eastern Software Solution</p>
    <p>✨ NEW: Sample Template Mode & Video Analysis | Built with Streamlit</p>
</div>

""", unsafe_allow_html=True)








